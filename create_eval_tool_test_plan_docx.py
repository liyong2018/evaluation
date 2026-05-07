from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "评估工具测试计划.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DEE7", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_fixed_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[idx])


def style_cell_text(cell, bold=False, color=INK, size=10.5):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, rows, widths, header=True):
    table = doc.add_table(rows=len(rows), cols=len(widths))
    set_fixed_table(table, widths)
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = str(value)
            if header and r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                style_cell_text(cell, bold=True, color=INK, size=10)
            else:
                style_cell_text(cell, color=INK, size=10)
    doc.add_paragraph()
    return table


def add_para(doc, text="", style=None, after=6, before=0, bold=False, color=INK, size=11):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6 if level == 2 else 4)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12,
                 bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_fixed_table(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in (
        ("Title", 23, RGBColor(0, 0, 0)),
        ("Subtitle", 14, GRAY),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True if style_name.startswith("Heading") else None


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("评估工具测试计划")
    set_run_font(hr, size=9, color=GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("内部测试资料")
    set_run_font(fr, size=9, color=GRAY)


def add_cover(doc):
    add_para(doc, "", after=20)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("评估工具测试计划")
    set_run_font(r, size=24, bold=True, color=RGBColor(0, 0, 0))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    sr = subtitle.add_run("系统计算结果与 Excel 人工计算结果一致性验证")
    set_run_font(sr, size=14, color=GRAY)

    rows = [
        ("文档版本", "V1.0"),
        ("适用对象", "评估工具计算准确性测试、验收测试及回归测试"),
        ("测试目标", "验证系统产出的评估结果与 Excel 严格按公式计算结果是否一致"),
        ("编制日期", date.today().strftime("%Y年%m月%d日")),
        ("编制单位", "测试单位 / 承建单位"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_fixed_table(table, [1800, 7560])
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        set_cell_shading(table.cell(i, 0), LIGHT_GRAY)
        style_cell_text(table.cell(i, 0), bold=True)
        style_cell_text(table.cell(i, 1))

    doc.add_page_break()


def build_doc():
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_cover(doc)

    add_heading(doc, "一、测试背景", 1)
    add_para(doc, "为验证评估工具在实际业务数据下的计算准确性、结果一致性和过程可追溯性，需对系统自动计算结果与人工 Excel 计算结果开展一致性比对测试。测试通过选取具有代表性的历史数据或回流数据作为评估样本，分别由承建单位使用系统自动计算、测试单位使用 Excel 按公式独立核算，并对两类结果逐项核验。")
    add_para(doc, "本测试计划用于规范测试数据选取、计算过程、结果比对、差异分析、问题整改及最终报告编制工作，确保测试结论客观、可复核、可追溯。")
    add_callout(doc, "核心验证口径", "系统计算和 Excel 计算必须使用同一批原始数据、同一套指标公式、同一套权重规则、同一套精度和四舍五入规则。")

    add_heading(doc, "二、测试目标", 1)
    for item in [
        "验证评估工具中各项指标计算公式是否正确实现。",
        "验证系统对基础数据的读取、字段映射、数据转换、缺失值处理和异常值处理是否符合规则。",
        "验证系统计算结果与 Excel 人工计算结果在乡镇、区县、市州及汇总层级是否一致。",
        "验证单项指标得分、权重得分、综合得分、排名、评价等级或评价结论是否准确。",
        "验证系统结果导出文件是否完整、规范，并能够支撑后续复核和归档。",
        "对发现的差异进行定位、分类、整改和复测，为系统上线、验收或正式应用提供依据。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "三、测试范围", 1)
    scope_rows = [
        ("测试模块", "测试内容", "关注重点"),
        ("数据导入", "导入评估基础数据、行政区划数据、指标数据", "数据条数、字段映射、数据类型、行政区划匹配"),
        ("指标计算", "执行单项指标、标准化、权重和综合得分计算", "公式一致性、精度规则、边界值处理"),
        ("汇总计算", "按乡镇、区县、市州、省级等层级汇总", "分组口径、汇总范围、层级关系"),
        ("结果导出", "导出系统计算结果和中间结果", "字段完整性、格式规范性、可追溯性"),
        ("结果比对", "与 Excel 人工计算结果逐项比对", "差异值、差异比例、影响范围"),
        ("问题整改", "记录、分析、修复并复测差异问题", "原因明确、闭环管理、复测通过"),
    ]
    add_table(doc, scope_rows, [1700, 3900, 3760])

    add_heading(doc, "四、测试依据", 1)
    for item in [
        "评估指标体系、指标解释及业务规则说明。",
        "各指标计算公式、权重规则、评分规则、排名规则及等级划分规则。",
        "系统需求规格说明书、概要设计说明书或详细设计说明书。",
        "数据字典、字段说明、行政区划编码规则及数据采集口径。",
        "测试双方确认的 Excel 计算模板、公式说明和人工核算规则。",
        "测试双方确认的精度规则、四舍五入规则和误差容忍范围。",
        "项目验收要求、测试管理要求及相关制度文件。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "五、测试数据", 1)
    add_heading(doc, "5.1 数据来源", 2)
    for item in [
        "选取 2020 年全量数据作为评估数据源。",
        "或选取 2024 年、2025 年部分回流数据作为评估数据源。",
        "若采用部分回流数据，则按照市州乡镇数量进行随机抽样，每个市州抽取约 30% 的乡镇参与测试。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.2 抽样原则", 2)
    for item in [
        "覆盖全部或主要市州，避免测试结果仅反映单一区域情况。",
        "覆盖不同数量规模的区县和乡镇，兼顾大样本区域和小样本区域。",
        "覆盖指标值正常、偏高、偏低、缺失、为 0、异常波动等不同数据特征。",
        "抽样过程应保留随机规则、抽样比例、抽样清单和确认记录。",
        "抽样结果应经测试单位和承建单位共同确认后固化使用。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.3 数据质量要求", 2)
    data_rows = [
        ("检查项", "要求", "处理方式"),
        ("数据版本", "系统计算和 Excel 计算使用同一批固化数据", "形成数据版本号或文件哈希记录"),
        ("字段完整性", "关键字段不得缺失，指标字段应与数据字典一致", "缺失字段需补充或说明原因"),
        ("行政区划", "市州、区县、乡镇名称及编码应完整准确", "编码异常需在计算前修正并留痕"),
        ("异常值", "异常值应提前识别并明确是否纳入计算", "按确认规则处理，不得单方修改"),
        ("重复数据", "同一地区、同一年份、同一指标不得重复入库", "重复记录需去重或确认取值规则"),
        ("单位和精度", "指标单位、百分比口径、小数位数保持一致", "统一转换后再进入计算"),
    ]
    add_table(doc, data_rows, [1900, 4200, 3260])

    add_heading(doc, "六、测试方法", 1)
    add_para(doc, "本次测试采用“同源数据、双轨计算、逐项比对、差异分析、整改复测”的方法开展。承建单位通过评估系统自动计算，测试单位通过 Excel 按照确认公式独立计算，双方计算完成后对结果进行逐项核验。")
    method_rows = [
        ("方法环节", "说明"),
        ("同源数据", "双方使用完全一致的原始数据源，不得分别修改数据。"),
        ("双轨计算", "系统自动计算与 Excel 人工计算独立进行，避免相互引用结果。"),
        ("逐项比对", "按地区、年度、指标、得分、排名和评价等级等维度进行比对。"),
        ("差异分析", "对不一致结果分类定位，明确是否为系统问题、Excel 公式问题、数据口径问题或精度问题。"),
        ("整改复测", "对确认问题完成修复后，使用同一数据源重新计算并复核。"),
    ]
    add_table(doc, method_rows, [2000, 7360])

    add_heading(doc, "七、测试流程", 1)
    flow_steps = [
        "测试准备：明确测试目标、范围、数据来源、系统环境、Excel 模板、公式口径、精度规则和责任分工。",
        "数据选取：选择 2020 年全量数据，或选择 2024 年、2025 年部分回流数据；部分回流数据按市州乡镇数量随机抽取约 30%。",
        "数据固化：形成测试数据清单，确认数据文件、数据批次、字段说明和抽样记录，作为双方唯一计算数据源。",
        "系统计算：承建单位导入测试数据，检查导入结果，执行系统评估计算并导出系统结果。",
        "Excel 计算：测试单位使用同一数据源，在 Excel 中严格按照评估公式和权重规则完成独立核算。",
        "结果提交：当日核算完成后，测试单位将 Excel 计算结果提交承建单位。",
        "结果比对：承建单位对系统结果和 Excel 结果进行字段对齐、逐项比对和差异标记。",
        "差异分析：对差异项记录差异值、涉及地区、涉及指标、影响层级和可能原因。",
        "整改复测：属于系统问题的，由承建单位整改后重新计算，并对差异项开展回归测试。",
        "报告编制：承建单位形成测试报告，测试单位参与复核并确认最终结论。",
    ]
    for step in flow_steps:
        add_number(doc, step)

    add_heading(doc, "八、比对规则", 1)
    add_heading(doc, "8.1 比对维度", 2)
    for item in ["年度维度", "市州维度", "区县维度", "乡镇维度", "指标维度", "综合得分维度", "排名及评价等级维度"]:
        add_bullet(doc, item)
    add_heading(doc, "8.2 比对内容", 2)
    compare_rows = [
        ("比对对象", "比对内容", "一致性要求"),
        ("原始数据", "指标原值、地区编码、年度、基础字段", "完全一致"),
        ("中间结果", "标准化值、换算值、加权值、分项得分", "按精度规则一致"),
        ("最终结果", "综合得分、排名、等级、评价结论", "不得出现实质差异"),
        ("汇总结果", "区县、市州、省级汇总值", "统计口径一致"),
        ("导出结果", "字段、格式、排序、文件完整性", "满足归档和复核要求"),
    ]
    add_table(doc, compare_rows, [1700, 4500, 3160])
    add_heading(doc, "8.3 精度和误差规则", 2)
    for item in [
        "原始数据保持数据源原始精度，不得随意截断或改写。",
        "中间计算结果建议保留不少于 4 位小数，用于差异追溯。",
        "最终得分按业务要求保留 2 位或 4 位小数。",
        "因四舍五入导致的微小差异，可在双方确认的误差范围内视为一致。",
        "评价等级、排名、结论类结果不得因精度处理产生不一致。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "九、测试环境", 1)
    env_rows = [
        ("环境项", "内容", "记录要求"),
        ("系统环境", "评估工具测试环境或验收环境", "记录访问地址、部署时间、系统版本"),
        ("数据库环境", "测试数据库或 H2/专用测试库", "记录数据库类型、数据批次、初始化脚本"),
        ("Excel 环境", "测试单位使用的 Excel 版本及计算模板", "记录模板版本、公式版本、操作人员"),
        ("浏览器环境", "如涉及前端操作，记录浏览器类型及版本", "记录截图或操作日志"),
        ("文件环境", "测试数据、导出结果和比对表存放目录", "记录文件名、生成时间和版本号"),
    ]
    add_table(doc, env_rows, [1700, 3900, 3760])

    add_heading(doc, "十、职责分工", 1)
    role_rows = [
        ("角色", "主要职责"),
        ("承建单位", "维护系统测试环境；导入数据并执行系统计算；导出系统结果；开展结果比对；分析系统差异；修复问题；编制测试报告。"),
        ("测试单位", "确认测试数据范围；使用 Excel 独立核算；复核 Excel 公式和结果；当日提交核算结果；参与差异复核和测试结论确认。"),
        ("双方共同", "确认数据口径、公式规则、精度规则、误差范围、差异原因、整改结果和最终测试结论。"),
    ]
    add_table(doc, role_rows, [1700, 7660])

    add_heading(doc, "十一、通过标准", 1)
    for item in [
        "系统导入数据与固化测试数据源一致，无漏数、错数、重复或字段映射错误。",
        "系统计算公式、权重规则、汇总规则与双方确认的 Excel 公式一致。",
        "系统结果与 Excel 计算结果整体一致，核心指标、综合得分、排名和评价等级无实质差异。",
        "差异项均已完成原因分析，并确认不影响最终评价结论，或已完成整改并通过复测。",
        "系统导出结果完整、准确、格式规范，能够支撑归档、追溯和验收。",
        "测试过程、测试数据、计算结果、比对结果和问题闭环记录完整。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "十二、不通过判定条件", 1)
    for item in [
        "系统计算结果与 Excel 计算结果存在较大差异，且无法合理解释。",
        "关键指标公式、权重规则、汇总规则或评价等级规则实现错误。",
        "综合得分、排名、评价等级或评价结论与 Excel 结果不一致。",
        "系统数据导入存在漏数、错数、字段映射错误或行政区划匹配错误。",
        "系统对缺失值、异常值、0 值或空值的处理规则与确认规则不一致。",
        "差异问题未完成定位、整改或复测，无法支撑测试通过结论。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "十三、问题管理与复测要求", 1)
    issue_rows = [
        ("字段", "填写要求"),
        ("问题编号", "按日期或批次统一编号，例如 ISSUE-20260506-001。"),
        ("问题描述", "说明差异现象、涉及地区、涉及指标和差异值。"),
        ("问题类型", "数据源差异、字段映射差异、公式差异、精度差异、汇总口径差异、系统缺陷等。"),
        ("严重程度", "按高、中、低划分；影响最终评价结论的应列为高。"),
        ("原因分析", "说明问题产生原因及责任归属。"),
        ("整改措施", "说明系统修复、数据修正、公式调整或规则确认情况。"),
        ("复测结论", "记录复测时间、复测人员、复测结果和是否关闭。"),
    ]
    add_table(doc, issue_rows, [1900, 7460])

    add_heading(doc, "十四、测试输出物", 1)
    output_rows = [
        ("序号", "输出物", "说明"),
        ("1", "测试数据清单", "记录数据年份、区域范围、抽样比例、样本清单和数据版本。"),
        ("2", "系统计算结果表", "承建单位从系统导出的计算结果。"),
        ("3", "Excel 人工计算结果表", "测试单位按公式独立计算形成的结果。"),
        ("4", "系统与 Excel 结果比对表", "记录一致项、差异项、差异值和差异比例。"),
        ("5", "差异问题记录表", "记录问题描述、原因、责任、整改和复测情况。"),
        ("6", "问题整改及复测记录", "记录系统修复情况和回归测试结果。"),
        ("7", "评估工具测试报告", "汇总测试过程、结果、问题和最终结论。"),
    ]
    add_table(doc, output_rows, [900, 2600, 5860])

    add_heading(doc, "十五、风险及应对措施", 1)
    risk_rows = [
        ("风险", "影响", "应对措施"),
        ("数据口径不一致", "导致系统和 Excel 结果无法有效比对", "测试前统一确认数据来源、字段含义、单位、统计口径和数据版本。"),
        ("Excel 公式引用错误", "人工核算结果失真", "测试单位复核模板公式，关键公式由双方共同确认。"),
        ("精度规则不一致", "产生大量小数差异", "提前明确小数位数、四舍五入规则和误差容忍范围。"),
        ("抽样代表性不足", "测试结论覆盖性不足", "覆盖不同市州、不同规模乡镇和不同数据特征。"),
        ("差异原因定位困难", "延误整改和验收", "保留系统中间结果和 Excel 中间结果，支持逐级追溯。"),
        ("测试时间不足", "影响核算质量", "提前准备数据、模板和比对工具，优先完成核心指标比对。"),
    ]
    add_table(doc, risk_rows, [2200, 3000, 4160])

    add_heading(doc, "十六、测试计划安排", 1)
    schedule_rows = [
        ("阶段", "主要工作", "责任方", "输出物"),
        ("准备阶段", "确认测试范围、公式规则、测试环境、数据来源和 Excel 模板", "双方", "测试准备确认记录"),
        ("数据阶段", "抽取或确认测试数据，形成测试数据清单并固化版本", "双方", "测试数据清单"),
        ("计算阶段", "系统自动计算和 Excel 人工独立计算", "承建单位 / 测试单位", "系统结果表、Excel 结果表"),
        ("比对阶段", "开展结果比对、标记差异并进行原因分析", "承建单位", "结果比对表、差异问题表"),
        ("整改阶段", "修复系统问题并开展复测", "承建单位 / 测试单位", "整改及复测记录"),
        ("总结阶段", "编制并确认测试报告", "承建单位 / 双方", "评估工具测试报告"),
    ]
    add_table(doc, schedule_rows, [1400, 4300, 1700, 1960])

    add_heading(doc, "十七、测试结论要求", 1)
    conclusion_rows = [
        ("结论类型", "判定说明"),
        ("通过", "系统计算结果与 Excel 计算结果一致，满足测试目标，可支撑后续验收或正式使用。"),
        ("有条件通过", "存在少量非关键差异，差异原因明确且不影响最终评价结论，后续按记录优化。"),
        ("不通过", "存在关键计算差异、评价结果差异或系统计算逻辑问题，需整改后重新测试。"),
        ("待复测", "问题已完成整改，但尚需基于相同数据或补充数据进行复测确认。"),
    ]
    add_table(doc, conclusion_rows, [1800, 7560])

    add_heading(doc, "十八、附件模板", 1)
    add_heading(doc, "附件 1：结果比对表模板", 2)
    compare_template = [
        ("序号", "年度", "市州", "区县", "乡镇", "指标名称", "系统结果", "Excel 结果", "差异值", "是否一致"),
        ("1", "", "", "", "", "", "", "", "", ""),
        ("2", "", "", "", "", "", "", "", "", ""),
    ]
    add_table(doc, compare_template, [650, 800, 900, 900, 1000, 1500, 900, 900, 800, 1010])
    add_heading(doc, "附件 2：差异问题记录表模板", 2)
    issue_template = [
        ("问题编号", "差异描述", "问题类型", "严重程度", "原因分析", "整改措施", "复测结论"),
        ("", "", "", "", "", "", ""),
        ("", "", "", "", "", "", ""),
    ]
    add_table(doc, issue_template, [1300, 2200, 1300, 1000, 1500, 1400, 660])

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
